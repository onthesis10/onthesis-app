# /app/utils.py

import os
import requests
import json
import re
from io import BytesIO

# Import untuk AI
import google.generativeai as genai

# Import untuk Ekspor Dokumen
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import black
from bs4 import BeautifulSoup # Library untuk parsing HTML, tambahkan 'beautifulsoup4' ke requirements.txt

# --- Inisialisasi Model AI ---
# Mengambil kunci API dari environment variable
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
model = None
if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        # Menggunakan model yang lebih baru dan sesuai untuk tugas kompleks
        model = genai.GenerativeModel('gemini-1.5-flash')
        print("Model Gemini (gemini-1.5-flash) berhasil diinisialisasi dalam utils.py")
    except Exception as e:
        print(f"Error saat menginisialisasi Gemini di utils.py: {e}")
else:
    print("Peringatan: GEMINI_API_KEY tidak ditemukan. Fungsi AI tidak akan bekerja.")

# --- Fungsi Pencarian Referensi ---
def search_references_crossref(keywords, rows=10):
    """
    Mencari referensi dari CrossRef API berdasarkan kata kunci.
    Mengembalikan daftar kamus (dictionaries) yang berisi detail referensi.
    """
    if not keywords:
        return []
    
    try:
        # Menggunakan email untuk 'mailto' adalah praktik yang baik saat menggunakan API CrossRef
        headers = {'mailto': os.getenv('USER_EMAIL', 'default.email@example.com')}
        params = {
            "query.bibliographic": keywords,
            "rows": rows,
            "sort": "relevance"
        }
        response = requests.get("https://api.crossref.org/works", params=params, headers=headers)
        response.raise_for_status()
        data = response.json()
        
        references = []
        for item in data.get('message', {}).get('items', []):
            title = (item.get('title') or ['N/A'])[0]
            authors_list = item.get('author', [])
            authors = ", ".join(
                filter(None, [f"{author.get('given', '')} {author.get('family', '')}".strip() for author in authors_list])
            )
            year = item.get('created', {}).get('date-parts', [[None]])[0][0]
            doi = item.get('DOI', 'N/A')
            container_title = (item.get('container-title') or ['N/A'])[0]

            references.append({
                "title": title,
                "authors": authors,
                "year": year,
                "doi": doi,
                "journal": container_title
            })
        return references
    except requests.exceptions.RequestException as e:
        print(f"Error saat request ke CrossRef: {e}")
        return []
    except Exception as e:
        print(f"Error saat memproses data CrossRef: {e}")
        return []

# --- Fungsi-fungsi Terkait AI ---
def _call_gemini_with_retry(prompt, is_json_output=False, retries=3, delay=5):
    """Fungsi internal untuk memanggil Gemini dengan penanganan error dan retry."""
    if not model:
        raise ConnectionError("Model AI tidak terinisialisasi. Periksa GEMINI_API_KEY.")
    
    # Konfigurasi khusus jika output yang diharapkan adalah JSON
    generation_config = genai.types.GenerationConfig(
        response_mime_type="application/json"
    ) if is_json_output else None

    for attempt in range(retries):
        try:
            response = model.generate_content(prompt, generation_config=generation_config)
            # Membersihkan output jika JSON
            if is_json_output:
                # Menghapus markdown backticks dan membersihkan whitespace
                cleaned_text = re.sub(r'```json\s*|\s*```', '', response.text.strip())
                return json.loads(cleaned_text)
            return response.text
        except Exception as e:
            print(f"Attempt {attempt + 1} gagal: {e}")
            if attempt < retries - 1:
                time.sleep(delay)
            else:
                raise e # Lemparkan error setelah percobaan terakhir

def generate_outline_with_ai(research_inputs):
    """
    Membuat kerangka (outline) Bab 2 yang terstruktur dalam format JSON.
    """
    prompt = f"""
    Sebagai seorang ahli metodologi penelitian, buatlah kerangka (outline) yang komprehensif untuk BAB II: KAJIAN PUSTAKA.

    Judul Penelitian: "{research_inputs['title']}"
    Kata Kunci Utama: "{research_inputs['mainKeywords']}"
    Bahasa Penulisan: {research_inputs['language']}

    Struktur outline harus mencakup, namun tidak terbatas pada:
    1.  Landasan Teori: Uraikan teori-teori utama yang relevan.
    2.  Penelitian Terdahulu yang Relevan: Tinjauan studi sebelumnya.
    3.  Kerangka Pemikiran: Sintesis dari teori dan penelitian terdahulu.
    4.  Hipotesis Penelitian (jika relevan).

    HASIL HARUS DALAM FORMAT JSON yang valid, berupa sebuah list dari objects. Setiap object harus memiliki key "title" untuk judul sub-bab.
    Contoh format:
    [
        {{"title": "2.1 Landasan Teori"}},
        {{"title": "2.1.1 Teori A"}},
        {{"title": "2.1.2 Teori B"}},
        {{"title": "2.2 Penelitian Terdahulu"}}
    ]
    """
    try:
        return _call_gemini_with_retry(prompt, is_json_output=True)
    except Exception as e:
        print(f"Gagal membuat outline dengan AI, menggunakan fallback: {e}")
        # Jika gagal, kembalikan outline default
        return [
            {"title": "2.1 Landasan Teori"},
            {"title": "2.2 Penelitian Terdahulu"},
            {"title": "2.3 Kerangka Pemikiran"},
            {"title": "2.4 Hipotesis Penelitian"}
        ]

def generate_subchapter_with_ai(research_inputs, subchapter, references):
    """
    Menghasilkan konten mendalam untuk satu sub-bab menggunakan AI.
    """
    references_str = "\n".join([f"- {ref['title']} oleh {ref['authors']} ({ref['year']})" for ref in references[:5]])
    
    prompt = f"""
    Anda adalah seorang penulis akademik. Tugas Anda adalah menulis konten untuk sub-bab: "{subchapter['title']}" dalam konteks penelitian berjudul "{research_inputs['title']}".

    Gaya Penulisan: {research_inputs['language']}, formal, dan objektif.
    Target Panjang: Sekitar 500-700 kata.

    Gunakan referensi berikut sebagai inspirasi dan dasar argumen. Lakukan parafrase dan sintesis, JANGAN PLAGIAT:
    {references_str}

    Instruksi:
    1.  Jelaskan konsep, teori, atau temuan utama yang relevan dengan sub-bab ini.
    2.  Kaitkan pembahasan secara langsung dengan variabel atau konteks penelitian.
    3.  Gunakan sitasi dalam format (Nama, Tahun) jika Anda mengutip ide dari referensi.
    4.  Output harus dalam format Markdown yang rapi.
    """
    try:
        return _call_gemini_with_retry(prompt)
    except Exception as e:
        return f"### {subchapter['title']}\n\n*Error: Gagal menghasilkan konten untuk sub-bab ini. {e}*"

def review_and_compile_with_ai(research_inputs, chapters, references):
    """
    Menggabungkan semua bab, melakukan review, dan membuat daftar pustaka.
    """
    full_content = "\n\n".join(chapters.values())

    review_prompt = f"""
    Sebagai seorang editor ahli, review draf Bab 2 berikut untuk penelitian berjudul "{research_inputs['title']}".

    Draf Awal:
    ---
    {full_content}
    ---

    Tugas Anda:
    1.  Periksa koherensi dan alur logika antar sub-bab.
    2.  Tambahkan kalimat transisi yang mulus untuk menghubungkan setiap bagian.
    3.  Perbaiki tata bahasa dan pastikan gaya penulisan konsisten ({research_inputs['language']}).
    4.  Pastikan semua klaim didukung oleh argumen yang logis.

    Kembalikan HANYA teks Bab 2 yang sudah direview dan disempurnakan dalam format Markdown.
    """
    try:
        final_content = _call_gemini_with_retry(review_prompt)
    except Exception as e:
        print(f"Gagal mereview dengan AI, menggunakan konten mentah: {e}")
        final_content = full_content

    # Membuat daftar pustaka dalam format HTML
    references_html = f"<h3>Daftar Pustaka (Format: {research_inputs.get('citationFormat', 'APA 7')})</h3>\n<ul>"
    for ref in references:
        # Logika sederhana untuk format APA, bisa diperluas
        references_html += f"<li>{ref.get('authors', 'N/A')}. ({ref.get('year', 'N/A')}). {ref.get('title', 'N/A')}. <em>{ref.get('journal', 'N/A')}</em>. https://doi.org/{ref.get('doi', '')}</li>\n"
    references_html += "</ul>"
    
    return final_content, references_html

# --- Fungsi-fungsi Ekspor Dokumen ---
def _parse_html_for_docx(soup, document):
    """Fungsi internal untuk mem-parsing HTML dan menambahkannya ke dokumen Word."""
    for tag in soup.find_all(['h3', 'p', 'ul', 'li']):
        if tag.name == 'h3':
            p = document.add_paragraph(style='Heading 3')
            p.add_run(tag.get_text()).bold = True
        elif tag.name == 'p':
            document.add_paragraph(tag.get_text())
        elif tag.name == 'ul':
            for li in tag.find_all('li'):
                document.add_paragraph(li.get_text(), style='List Bullet')

def export_to_docx(content_html, references_html):
    """
    Mengekspor konten dari HTML ke file DOCX dengan parsing dasar.
    """
    document = Document()
    document.add_heading('BAB II KAJIAN PUSTAKA', level=1)
    
    # Parsing konten utama
    soup_content = BeautifulSoup(content_html, 'html.parser')
    _parse_html_for_docx(soup_content, document)
    
    document.add_page_break()
    
    # Parsing daftar pustaka
    soup_refs = BeautifulSoup(references_html, 'html.parser')
    _parse_html_for_docx(soup_refs, document)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

def _parse_html_for_pdf(soup, story, styles):
    """Fungsi internal untuk mem-parsing HTML dan menambahkannya ke story ReportLab."""
    for tag in soup.find_all(['h3', 'p', 'ul', 'li']):
        if tag.name == 'h3':
            story.append(Paragraph(tag.get_text(), styles['h3']))
            story.append(Spacer(1, 0.1*inch))
        elif tag.name == 'p':
            story.append(Paragraph(tag.get_text(), styles['BodyText']))
            story.append(Spacer(1, 0.1*inch))
        elif tag.name == 'ul':
            for li in tag.find_all('li'):
                # Menambahkan bullet point secara manual
                p_text = f"• {li.get_text()}"
                story.append(Paragraph(p_text, styles['BodyText']))
            story.append(Spacer(1, 0.1*inch))

def export_to_pdf(content_html, references_html):
    """
    Mengekspor konten dari HTML ke file PDF menggunakan ReportLab.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    
    story = []
    story.append(Paragraph("BAB II KAJIAN PUSTAKA", styles['h1']))
    story.append(Spacer(1, 0.2*inch))

    # Parsing konten utama
    soup_content = BeautifulSoup(content_html, 'html.parser')
    _parse_html_for_pdf(soup_content, story, styles)
    
    story.append(PageBreak())
    
    # Parsing daftar pustaka
    soup_refs = BeautifulSoup(references_html, 'html.parser')
    _parse_html_for_pdf(soup_refs, story, styles)
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
