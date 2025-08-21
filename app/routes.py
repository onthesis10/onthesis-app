# ========================================================================
# File: app/routes.py
# Deskripsi: Versi LENGKAP dengan semua perbaikan bug untuk fitur ANOVA
#            dan fitur-fitur lainnya.
# Perubahan Final v2:
# 1. FIX TypeError: Mengonversi numpy.bool_ menjadi bool standar Python
#    untuk mencegah error "not JSON serializable".
# ========================================================================

# --- Impor Library ---
import os
import json
import re
import requests
import google.generativeai as genai
import time
import midtransclient
from datetime import date, datetime, timedelta
from werkzeug.utils import secure_filename
import uuid
import io
import base64
from concurrent.futures import ThreadPoolExecutor

# --- Impor untuk Analisis Statistik ---
from scipy import stats
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')  # non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns
# --- Impor tambahan untuk ANOVA Profesional ---
from statsmodels.stats.multicomp import pairwise_tukeyhsd
import statsmodels.api as sm
from statsmodels.formula.api import ols
import pingouin as pg


# --- Impor dari __init__.py ---
from app import app, db, login_manager

# Impor untuk framework Flask dan ekstensi
from flask import render_template, jsonify, request, redirect, url_for, flash, send_file
from flask_cors import CORS
from flask_login import (
    UserMixin,
    login_user,
    logout_user,
    login_required,
    current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
from firebase_admin import auth, firestore

# Impor untuk analisis dokumen
import PyPDF2
import docx

# --- Impor untuk Ekspor Dokumen ---
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from docx import Document
    from docx.shared import Inches
    PDF_EXPORT_ENABLED = True
    WORD_EXPORT_ENABLED = True
except ImportError:
    print("PERINGATAN: Library 'reportlab' atau 'python-docx' tidak terinstal. Fitur ekspor tidak akan berfungsi.")
    PDF_EXPORT_ENABLED = False
    WORD_EXPORT_ENABLED = False


# --- Konfigurasi Tambahan ---
try:
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
except Exception as e:
    print(f"Peringatan: Gagal mengkonfigurasi Gemini API. Error: {e}")

try:
    server_key = os.getenv('MIDTRANS_SERVER_KEY')
    client_key = os.getenv('MIDTRANS_CLIENT_KEY')
    midtrans_snap = midtransclient.Snap(
        is_production=False,
        server_key=server_key,
        client_key=client_key
    )
except Exception as e:
    print(f"Peringatan: Gagal mengkonfigurasi Midtrans. Error: {e}")
    midtrans_snap = None

OUTPUT_DIR = os.path.join(app.static_folder, 'outputs')
os.makedirs(OUTPUT_DIR, exist_ok=True)

sns.set_style('whitegrid')
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['Arial', 'DejaVu Sans']
plt.rcParams['axes.titleweight'] = 'bold'
plt.rcParams['axes.titlesize'] = 14
plt.rcParams['axes.labelsize'] = 12
plt.rcParams['xtick.labelsize'] = 10
plt.rcParams['ytick.labelsize'] = 10
plt.rcParams['figure.figsize'] = (8,5)


# =========================================================================
# FUNGSI HELPER
# =========================================================================
def read_pdf(file_stream):
    reader = PyPDF2.PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def read_docx(file_stream):
    doc = docx.Document(file_stream)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def create_plot_as_base64(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    plt.close(fig)
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode('utf-8')

def make_api_request_with_retry(url, headers, params=None, timeout=25, retries=3, backoff_factor=2):
    for attempt in range(retries):
        try:
            response = requests.get(url, headers=headers, params=params, timeout=timeout)
            if response.status_code == 404:
                print(f"Sumber tidak ditemukan (404) di URL: {url}. Melewati.")
                return None
            response.raise_for_status()
            return response
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 429:
                if attempt < retries - 1:
                    delay = backoff_factor ** attempt
                    print(f"Rate limit terdeteksi. Mencoba lagi dalam {delay} detik...")
                    time.sleep(delay)
                else:
                    print("Gagal setelah beberapa kali percobaan. Melemparkan error.")
                    raise
            else:
                raise
        except requests.exceptions.RequestException as e:
            print(f"Error koneksi: {e}")
            if attempt < retries - 1:
                delay = backoff_factor ** attempt
                time.sleep(delay)
            else:
                raise
    return None

def sanitize_nan(data):
    """Recursively converts NaN/inf values to None for JSON compatibility."""
    if isinstance(data, dict):
        return {k: sanitize_nan(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [sanitize_nan(i) for i in data]
    elif isinstance(data, float) and (np.isnan(data) or np.isinf(data)):
        return None
    return data

# =========================================================================
# MODEL PENGGUNA & LOADER
# =========================================================================
class User(UserMixin):
    def __init__(self, id, displayName, password_hash=None, email=None, picture=None, pro_expiry_date=None, legacy_is_pro=False):
        self.id = id
        self.displayName = displayName
        self.username = displayName
        self.email = email
        self.picture = picture
        self.password_hash = password_hash
        self.pro_expiry_date = pro_expiry_date
        self.legacy_is_pro = legacy_is_pro

    @property
    def is_pro(self):
        if self.pro_expiry_date and isinstance(self.pro_expiry_date, datetime):
            return self.pro_expiry_date > datetime.now()
        if self.legacy_is_pro and self.pro_expiry_date is None:
            return True
        return False

    def check_password(self, password):
        if self.password_hash is None: return False
        return check_password_hash(self.password_hash, password)

@login_manager.user_loader
def load_user(user_id):
    if not db: return None
    try:
        user_doc = db.collection('users').document(user_id).get()
        if user_doc.exists:
            user_data = user_doc.to_dict()
            return User(
                id=user_id, 
                displayName=user_data.get('displayName'), 
                email=user_data.get('email'), 
                password_hash=user_data.get('password_hash'), 
                picture=user_data.get('picture'),
                pro_expiry_date=user_data.get('proExpiryDate'),
                legacy_is_pro=user_data.get('isPro', False)
            )
        return None
    except Exception as e:
        print(f"Error saat memuat pengguna dari Firestore: {e}")
        return None

# =========================================================================
# FUNGSI HELPER UNTUK PEMBATASAN FITUR
# =========================================================================
def check_and_update_usage(user_id, feature_name):
    FEATURE_LIMITS = {
        'paraphrase': 5, 'chat': 10, 'search': 5, 'citation': 15
    }
    limit = FEATURE_LIMITS.get(feature_name)
    if limit is None: return True, "OK"
    user_ref = db.collection('users').document(user_id)
    user_doc = user_ref.get()
    if not user_doc.exists: return False, "Pengguna tidak ditemukan."
    today_str = date.today().isoformat()
    usage_data = user_doc.to_dict().get('usage_limits', {})
    last_reset = usage_data.get('last_reset_date')
    if last_reset != today_str:
        citation_total = usage_data.get('citation_count', 0)
        usage_data = {
            'paraphrase_count': 0, 'chat_count': 0, 'search_count': 0,
            'writing_assistant_count': 0, 'data_analysis_count': 0, 'export_doc_count': 0,
            'last_reset_date': today_str, 'citation_count': citation_total,
            'generate_theory_count': 0
        }
        user_ref.set({'usage_limits': usage_data}, merge=True)
    count_key = f"{feature_name}_count"
    current_count = usage_data.get(count_key, 0)
    if current_count >= limit:
        if feature_name == 'citation':
             return False, f"Anda telah mencapai batas total {limit} referensi untuk akun gratis."
        return False, f"Anda telah mencapai batas penggunaan harian ({limit}x) untuk fitur ini. Silakan upgrade ke PRO."
    user_ref.update({f'usage_limits.{count_key}': firestore.Increment(1)})
    return True, "OK"

def check_and_update_pro_trial(user_id, feature_name):
    PRO_TRIAL_LIMITS = {'writing_assistant': 3, 'data_analysis': 3, 'export_doc': 1, 'generate_theory': 2}
    limit = PRO_TRIAL_LIMITS.get(feature_name)
    if limit is None: return True, "OK"
    user_ref = db.collection('users').document(user_id)
    user_doc = user_ref.get()
    if not user_doc.exists: return False, "Pengguna tidak ditemukan."
    usage_data = user_doc.to_dict().get('usage_limits', {})
    count_key = f"{feature_name}_count"
    current_count = usage_data.get(count_key, 0)
    if current_count >= limit:
        return False, "UPGRADE_REQUIRED"
    user_ref.update({f'usage_limits.{count_key}': firestore.Increment(1)})
    return True, "OK"

# =========================================================================
# RUTE-RUTE OTENTIKASI DAN HALAMAN
# =========================================================================
@app.route('/login', methods=['GET'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    firebase_config = { 
        "apiKey": os.getenv("FIREBASE_API_KEY"), 
        "authDomain": os.getenv("FIREBASE_AUTH_DOMAIN"), 
        "projectId": os.getenv("FIREBASE_PROJECT_ID"), 
        "storageBucket": os.getenv("FIREBASE_STORAGE_BUCKET"), 
        "messagingSenderId": os.getenv("FIREBASE_MESSAGING_SENDER_ID"), 
        "appId": os.getenv("FIREBASE_APP_ID"), 
        "measurementId": os.getenv("FIREBASE_MEASUREMENT_ID") 
    }
    firebase_config_filtered = {k: v for k, v in firebase_config.items() if v is not None}
    return render_template('login.html', firebase_config=firebase_config_filtered)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Anda telah berhasil logout.', 'success')
    return redirect(url_for('login'))

@app.route('/')
@app.route('/dashboard')
@login_required
def dashboard(): return render_template('dashboard.html')

@app.route('/projects')
@login_required
def projects(): return render_template('projects.html')

@app.route('/search-references')
@login_required
def search_references(): return render_template('search_references.html')

@app.route('/citation-management')
@login_required
def citation_management(): return render_template('citation_management.html')

@app.route('/paraphrase-ai')
@login_required
def paraphrase_ai(): return render_template('paraphrase_ai.html')

@app.route('/chat-ai')
@login_required
def chat_ai(): return render_template('chat_ai.html')

@app.route('/writing-assistant')
@login_required
def writing_assistant(): 
    return render_template('writing_assistant.html')

@app.route('/generator-latar-belakang')
@login_required
def generator_latar_belakang():
    return render_template('generator_latar_belakang.html')

@app.route('/generator-rumusan-masalah')
@login_required
def generator_rumusan_masalah():
    return render_template('generator_rumusan_masalah.html')

@app.route('/generator-kajian-teori')
@login_required
def generator_kajian_teori():
    return render_template('generator_kajian_teori.html')

@app.route('/data-analysis')
@login_required
def data_analysis(): return render_template('data_analysis.html')

@app.route('/normality-test')
@login_required
def normality_test(): return render_template('normality_test.html')

@app.route('/homogeneity_test')
@login_required
def homogeneity_test(): return render_template('homogeneity_test.html')

@app.route('/t-test')
@login_required
def t_test():
    return render_template('t_test.html')

@app.route('/anova-test')
@login_required
def anova_test():
    return render_template('anova_test.html')

@app.route('/descriptive_statistics')
@login_required
def descriptive_statistics():
    return render_template('descriptive_statistics.html')

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def user_profile():
    if request.method == 'POST':
        try:
            new_name = request.form.get('name')
            if not new_name or len(new_name) < 3:
                flash('Nama tampilan harus memiliki setidaknya 3 karakter.', 'danger')
                return redirect(url_for('user_profile'))
            user_id = current_user.id
            db.collection('users').document(user_id).update({'displayName': new_name})
            auth.update_user(user_id, display_name=new_name)
            flash('Profil berhasil diperbarui!', 'success')
        except Exception as e:
            flash(f'Terjadi kesalahan saat memperbarui profil: {e}', 'danger')
        return redirect(url_for('user_profile'))
    client_key = os.getenv('MIDTRANS_CLIENT_KEY')
    return render_template('user-profile.html', midtrans_client_key=client_key)

@app.route('/upgrade')
@login_required
def upgrade_page():
    client_key = os.getenv('MIDTRANS_CLIENT_KEY')
    return render_template('upgrade.html', client_key=client_key)

# =========================================================================
# RUTE API BARU UNTUK OTENTIKASI & LAINNYA
# =========================================================================
@app.route('/api/verify-google-token', methods=['POST'])
def verify_google_token():
    try:
        token = request.json['token']
        decoded_token = auth.verify_id_token(token)
        uid = decoded_token['uid']
        user_ref = db.collection('users').document(uid)
        user_doc = user_ref.get()

        if user_doc.exists:
            user = load_user(uid)
        else:
            user_data = {
                'displayName': decoded_token.get('name', decoded_token.get('email')),
                'email': decoded_token.get('email'),
                'picture': decoded_token.get('picture'),
                'isPro': False, 
                'password_hash': None,
                'proExpiryDate': None
            }
            user_ref.set(user_data)
            user = load_user(uid)
        
        login_user(user)
        return jsonify({'status': 'success', 'redirect_url': url_for('dashboard')})
    except Exception as e:
        print(f"Error verifikasi token: {e}")
        return jsonify({'status': 'error', 'message': f'Verifikasi token gagal: {e}'}), 401

@app.route('/api/verify-email-token', methods=['POST'])
def verify_email_token():
    try:
        token = request.json.get('token')
        if not token:
            return jsonify({'status': 'error', 'message': 'Token tidak ditemukan.'}), 400

        decoded_token = auth.verify_id_token(token)
        uid = decoded_token['uid']
        
        user = load_user(uid)
        if user:
            login_user(user)
            return jsonify({'status': 'success', 'redirect_url': url_for('dashboard')})
        else:
            return jsonify({'status': 'error', 'message': 'Pengguna tidak ditemukan di database.'}), 404
            
    except auth.InvalidIdTokenError:
        return jsonify({'status': 'error', 'message': 'Token tidak valid atau kedaluwarsa.'}), 401
    except Exception as e:
        print(f"Error verifikasi token email: {e}")
        return jsonify({'status': 'error', 'message': 'Terjadi kesalahan di server.'}), 500


@app.route('/api/check-pro-trial-usage', methods=['POST'])
@login_required
def check_pro_trial_usage():
    if current_user.is_pro:
        return jsonify({'allowed': True})
    
    try:
        data = request.get_json()
        feature_name = data.get('feature')
        if not feature_name:
            return jsonify({'error': 'Nama fitur diperlukan.'}), 400
        
        is_allowed, message = check_and_update_pro_trial(current_user.id, feature_name)
        
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({
                    'allowed': False, 
                    'redirect': url_for('upgrade_page')
                }), 429
            return jsonify({'allowed': False, 'message': message}), 429
        
        return jsonify({'allowed': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/writing-assistant', methods=['POST'])
@login_required
def api_writing_assistant():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'writing_assistant')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429
    try:
        data = request.get_json()
        task = data.get('task')
        context = data.get('context')
        if not task or not context: return jsonify({'error': 'Task dan context diperlukan.'}), 400
        
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = ""

        if task == 'generate_outline':
            prompt = f"Buatkan kerangka skripsi yang terstruktur dan logis berdasarkan judul berikut: \"{context}\""
        
        elif task == 'generate_abstract':
            prompt = f"Buatkan draf abstrak yang ringkas dan padat (sekitar 200-250 kata) berdasarkan isi skripsi berikut:\n\n{context}"
        
        elif task == 'generate_verified_background':
            if not isinstance(context, dict):
                return jsonify({'error': 'Context untuk generate_background harus berupa objek.'}), 400
            
            topic = context.get('topic', '')
            major = context.get('major', '')
            year = context.get('year', '')
            citation_style = context.get('citationStyle', 'APA 7')
            paragraph_count = context.get('paragraphCount', '4')

            references_text = ""
            try:
                search_query = f"{topic} in {major}"
                crossref_url = 'https://api.crossref.org/works'
                params = {'query.bibliographic': search_query, 'rows': 5, 'sort': 'relevance'}
                if year:
                    params['filter'] = f'from-pub-date:{year}-01-01,until-pub-date:{year}-12-31'
                
                headers = {'User-Agent': 'OnThesisApp/1.0 (mailto:contact@onthesis.app)'}
                crossref_response = requests.get(crossref_url, params=params, headers=headers, timeout=20)
                
                found_references = []
                if crossref_response.ok:
                    items = crossref_response.json().get('message', {}).get('items', [])
                    for item in items:
                        title = item.get('title', [''])[0]
                        authors_list = item.get('author', [])
                        authors = ", ".join([f"{author.get('family', '')}, {author.get('given', '')[0]}." for author in authors_list if author.get('family') and author.get('given')])
                        pub_year = item.get('issued', {}).get('date-parts', [[None]])[0][0]
                        journal = item.get('container-title', [''])[0]
                        doi = item.get('DOI', '')
                        
                        if title and authors and pub_year:
                            ref_info = f"Judul: {title}, Penulis: {authors}, Tahun: {pub_year}"
                            if journal: ref_info += f", Jurnal: {journal}"
                            if doi: ref_info += f", DOI: https://doi.org/{doi}"
                            found_references.append(ref_info)

                if found_references:
                    formatting_prompt = f"""
                    Berdasarkan daftar informasi referensi berikut, format masing-masing ke dalam gaya sitasi {citation_style}.
                    Sajikan hasilnya sebagai daftar bernomor. Pastikan formatnya benar dan konsisten.

                    Informasi Referensi:
                    {chr(10).join(f'- {ref}' for ref in found_references)}
                    """
                    formatting_response = model.generate_content(formatting_prompt)
                    references_text = formatting_response.text
                else:
                    references_text = "Tidak ada referensi relevan yang ditemukan secara otomatis. Silakan tambahkan secara manual."

            except Exception as e:
                print(f"Gagal mencari referensi nyata: {e}")
                references_text = "Terjadi kesalahan saat mencari referensi. Bagian ini bisa diisi manual."
            
            prompt = f"""
                Anda adalah seorang asisten penulis skripsi ahli. Tugas Anda adalah membuat draf Latar Belakang Masalah berdasarkan informasi berikut:

                1.  **Topik Utama:** {topic}
                2.  **Bidang/Jurusan:** {major}

                Instruksi Penulisan:
                * **Jumlah Paragraf:** Tulis draf dalam {paragraph_count} paragraf yang terstruktur.
                * **Struktur:** Gunakan struktur berikut: Pengenalan topik, Masalah aktual (didukung data), Relevansi dan gap penelitian, dan Penutup menuju rumusan masalah.
                * **Sitasi:** Sitasi beberapa sumber dari daftar pustaka yang disediakan di dalam teks jika relevan, contoh: (Nama Penulis, Tahun).
                * **Daftar Pustaka:** Setelah semua paragraf, buat bagian baru dengan judul `### Daftar Pustaka`. Gunakan daftar referensi yang sudah diformat berikut ini:\n{references_text}

                Tulis draf Latar Belakang Masalah dan Daftar Pustaka sekarang dalam format Markdown.
            """
        elif task == 'generate_problem_statement':
            if not isinstance(context, dict):
                return jsonify({'error': 'Context harus berupa objek.'}), 400
            
            topic = context.get('topic', '')
            background = context.get('background', '')
            point_count = context.get('pointCount', '3')

            prompt = f"""
            Anda adalah seorang metodolog penelitian ahli. Tugas Anda adalah merumuskan pertanyaan penelitian (rumusan masalah) yang tajam dan relevan berdasarkan informasi berikut:

            1.  **Judul/Topik Penelitian:** {topic}
            2.  **Konteks Latar Belakang (jika ada):** {background}

            **Instruksi:**
            - Identifikasi variabel-variabel kunci, populasi, dan konteks dari informasi yang diberikan.
            - Buatlah {point_count} poin pertanyaan penelitian yang spesifik, terukur, dapat dicapai, relevan, dan terikat waktu (SMART), jika memungkinkan.
            - Fokus pada pertanyaan "Bagaimana", "Apakah ada hubungan/pengaruh", "Seberapa besar", atau "Faktor-faktor apa saja".
            - Sajikan hasilnya sebagai daftar bernomor dalam format Markdown.
            """
        else:
            return jsonify({'error': 'Task tidak valid.'}), 400
            
        response = model.generate_content(prompt)
        return jsonify({'generated_text': response.text})
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# --- FUNGSI-FUNGSI PENCARIAN BARU ---
def search_core(keywords):
    print(f"Mencari di CORE dengan keywords: {keywords}")
    core_api_key = os.getenv('CORE_API_KEY')
    if not core_api_key: return []

    keyword_list = [keyword.strip() for keyword in keywords.split(',')]
    core_query = " AND ".join(keyword_list)
    
    core_url = f"https://api.core.ac.uk/v3/search/works?q={core_query}&limit=15"
    core_headers = {"Authorization": f"Bearer {core_api_key}"}
    
    core_response = make_api_request_with_retry(core_url, headers=core_headers)
    if not core_response: return []
    
    core_results = core_response.json().get('results', [])
    if not core_results: return []

    processed_references = []
    crossref_headers = {'User-Agent': 'OnThesisApp/1.0 (mailto:dev@onthesis.app)'}
    
    for item in core_results:
        doi = item.get('doi')
        if not doi: continue

        crossref_url = f"https://api.crossref.org/works/{doi}"
        crossref_response = make_api_request_with_retry(crossref_url, headers=crossref_headers, timeout=10)
        
        if crossref_response and crossref_response.status_code == 200:
            crossref_data = crossref_response.json().get('message', {})
            if crossref_data.get('abstract'):
                authors = crossref_data.get('author', [])
                if not authors: continue

                authors_str_list = [a.get('family', '') for a in authors if a.get('family')]
                if not authors_str_list: continue
                
                authors_str = " & ".join(authors_str_list[:2])
                if len(authors_str_list) > 2: authors_str += ", et al."
                
                year_parts = crossref_data.get('issued', {}).get('date-parts', [[None]])[0]
                year = year_parts[0] if year_parts and year_parts[0] else "n.d."

                processed_references.append({
                    "title": crossref_data.get('title', ['N/A'])[0],
                    "authors_str": authors_str,
                    "year": year,
                    "abstract": re.sub('<[^<]+?>', '', crossref_data.get('abstract')),
                    "doi": doi
                })
    return processed_references

def search_openalex(keywords):
    print(f"Mencari di OpenAlex dengan keywords: {keywords}")
    base_url = "https://api.openalex.org/works"
    search_query = keywords.replace(",", " ")
    params = {'search': search_query, 'per-page': 10}
    response = make_api_request_with_retry(base_url, headers={}, params=params)
    if not response: return []
    
    results = []
    for item in response.json().get('results', []):
        if not item.get('abstract_inverted_index'): continue
        
        authors = [author['author']['display_name'] for author in item.get('authorships', [])]
        year = item.get('publication_year')
        
        abstract = ""
        if item.get('abstract_inverted_index'):
            abstract_dict = item['abstract_inverted_index']
            sorted_words = sorted(abstract_dict.items(), key=lambda item: item[1][0])
            abstract = ' '.join(word for word, pos in sorted_words)

        results.append({
            "title": item.get('display_name', 'N/A'),
            "authors_str": ", ".join(authors[:2]) + (", et al." if len(authors) > 2 else ""),
            "year": year,
            "abstract": abstract,
            "doi": item.get('doi', '').replace('https://doi.org/', '')
        })
    return results

def search_doaj(keywords):
    print(f"Mencari di DOAJ dengan keywords: {keywords}")
    search_query = keywords.replace(",", "+")
    base_url = f"https://doaj.org/api/v2/search/articles/{search_query}"
    params = {'pageSize': 10}
    response = make_api_request_with_retry(base_url, headers={}, params=params)
    if not response: return []

    results = []
    for item in response.json().get('results', []):
        bibjson = item.get('bibjson', {})
        if not bibjson.get('abstract'): continue
        
        authors = [author['name'] for author in bibjson.get('author', [])]
        year = bibjson.get('year')
        doi = next((identifier['id'] for identifier in bibjson.get('identifier', []) if identifier.get('type') == 'doi'), None)

        results.append({
            "title": bibjson.get('title', 'N/A'),
            "authors_str": ", ".join(authors[:2]) + (", et al." if len(authors) > 2 else ""),
            "year": year,
            "abstract": bibjson.get('abstract'),
            "doi": doi
        })
    return results

def search_eric(keywords):
    print(f"Mencari di ERIC dengan keywords: {keywords}")
    base_url = "https://api.ies.ed.gov/eric/"
    params = {'search': keywords, 'rows': 10, 'format': 'json'}
    response = make_api_request_with_retry(base_url, headers={}, params=params)
    if not response: return []

    results = []
    for item in response.json().get('response', {}).get('docs', []):
        if not item.get('description'): continue
        
        authors = item.get('author', [])
        year = item.get('publicationdateyear')

        results.append({
            "title": item.get('title', 'N/A'),
            "authors_str": ", ".join(authors[:2]) + (", et al." if len(authors) > 2 else ""),
            "year": year,
            "abstract": item.get('description'),
            "doi": None
        })
    return results

def search_pubmed(keywords):
    print(f"Mencari di PubMed dengan keywords: {keywords}")
    api_key = os.getenv("PUBMED_API_KEY")
    if not api_key: return []
    
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    search_url = f"{base_url}esearch.fcgi"
    params = {'db': 'pubmed', 'term': keywords, 'retmax': 10, 'retmode': 'json', 'api_key': api_key}
    search_response = make_api_request_with_retry(search_url, headers={}, params=params)
    if not search_response: return []
    
    ids = search_response.json().get('esearchresult', {}).get('idlist', [])
    if not ids: return []

    summary_url = f"{base_url}esummary.fcgi"
    params = {'db': 'pubmed', 'id': ",".join(ids), 'retmode': 'json', 'api_key': api_key}
    summary_response = make_api_request_with_retry(summary_url, headers={}, params=params)
    if not summary_response: return []

    results = []
    for uid, data in summary_response.json().get('result', {}).items():
        if uid == 'uids': continue
        
        authors = [author['name'] for author in data.get('authors', [])]
        year = data.get('pubdate', '').split(' ')[0]
        doi = next((articleid['value'] for articleid in data.get('articleids', []) if articleid.get('idtype') == 'doi'), None)

        results.append({
            "title": data.get('title', 'N/A'),
            "authors_str": ", ".join(authors[:2]) + (", et al." if len(authors) > 2 else ""),
            "year": year,
            "abstract": f"Abstrak tidak tersedia langsung dari PubMed summary. Artikel membahas tentang {data.get('title', '')}.",
            "doi": doi
        })
    return results

# =========================================================================
# API BARU UNTUK GENERATOR KAJIAN TEORI (ALUR INTERAKTIF)
# =========================================================================

@app.route('/api/generate-outline-and-refs', methods=['POST'])
@login_required
def generate_outline_and_refs():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'generate_theory')
        if not is_allowed:
            return jsonify({'error': "Batas percobaan tercapai. Upgrade ke PRO."}), 429

    data = request.get_json()
    research_title = data.get('title', '')
    keywords = data.get('keywords', '')
    min_year = 2018

    if not research_title:
        return jsonify({"error": "Judul penelitian tidak boleh kosong."}), 400

    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt_outline = f"""
        Anda adalah seorang perencana penelitian ahli. Berdasarkan judul penelitian berikut, buatlah struktur Bab 2 (Kajian Teori) yang profesional.
        Judul: "{research_title}"
        Tugas Anda:
        1. Buat outline Bab 2 yang terdiri dari bagian utama seperti 'Landasan Teori', 'Penelitian Terdahulu', dan 'Kerangka Pemikiran'.
        2. Gunakan format penomoran huruf kapital untuk setiap sub-bab utama (Contoh: A. Konsep Media Sosial).
        3. Di bawah setiap sub-bab, sertakan array 'poin_pembahasan' yang berisi 3-4 poin kunci yang harus dijelaskan.
        4. Sertakan juga array 'kata_kunci_pencarian' yang relevan untuk setiap sub-bab.
        Berikan output HANYA dalam format JSON.
        Contoh:
        {{
          "outline": [
            {{
              "sub_bab": "A. Konsep Media Sosial",
              "poin_pembahasan": ["Definisi dan evolusi.", "Klasifikasi platform.", "Peran dalam komunikasi."],
              "kata_kunci_pencarian": "definisi media sosial, jenis platform media sosial"
            }}
          ]
        }}
        """
        outline_response = model.generate_content(prompt_outline)
        clean_json_string = re.sub(r'```json\s*|\s*```', '', outline_response.text.strip(), flags=re.DOTALL)
        research_plan = json.loads(clean_json_string).get('outline', [])
        if not research_plan:
            raise ValueError("AI gagal membuat outline yang valid.")

        all_references = []
        search_tasks = []
        with ThreadPoolExecutor(max_workers=15) as executor:
            for section in research_plan:
                search_keywords = section.get('kata_kunci_pencarian', '')
                if search_keywords:
                    search_tasks.append(executor.submit(search_core, search_keywords))
                    search_tasks.append(executor.submit(search_openalex, search_keywords))
                    search_tasks.append(executor.submit(search_doaj, search_keywords))
                    search_tasks.append(executor.submit(search_eric, search_keywords))
        
            for future in search_tasks:
                try: all_references.extend(future.result())
                except Exception as e: print(f"Gagal mengambil hasil pencarian: {e}")

        unique_references = []
        seen_titles = set()
        for ref in all_references:
            if not ref or not ref.get('title'): continue
            try:
                ref_year_str = str(ref.get('year', '0'))
                ref_year = int(re.search(r'\d{4}', ref_year_str).group()) if re.search(r'\d{4}', ref_year_str) else 0
                if min_year and ref_year < int(min_year): continue
            except: continue
            
            title_lower = ref['title'].lower()
            if title_lower not in seen_titles:
                unique_references.append(ref)
                seen_titles.add(title_lower)

        if len(unique_references) < 5:
            return jsonify({"error": f"Referensi yang ditemukan tidak cukup (hanya {len(unique_references)}). Coba dengan judul yang lebih umum."}), 404

        return jsonify({"outline": research_plan, "references": unique_references})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Terjadi kesalahan internal: {str(e)}"}), 500


@app.route('/api/generate-subchapter-content', methods=['POST'])
@login_required
def generate_subchapter_content():
    data = request.get_json()
    subchapter = data.get('subchapter')
    references = data.get('references', [])
    research_title = data.get('title', '')
    length_preference = data.get('length_preference', 'Normal')
    citation_style = data.get('citation_style', 'APA 7')

    if not subchapter or not references:
        return jsonify({"error": "Data sub-bab dan referensi diperlukan."}), 400

    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        processed_references = sorted(references, key=lambda x: x.get('year', 0), reverse=True)[:25]
        sources_text = ""
        for i, ref in enumerate(processed_references):
            ref['citation_placeholder'] = f"[{ref.get('authors_str', 'N/A').split(' ')[0].replace(',', '')}, {ref.get('year', 'n.d.')}]"
            sources_text += f"Sumber {i+1}:\n- Judul: {ref.get('title')}\n- Abstrak: {ref.get('abstract')}\n- Sitasi: {ref.get('citation_placeholder')}\n- DOI: {ref.get('doi')}\n\n"
        
        length_instruction = "Tulis pembahasan yang sangat mendalam dan komprehensif, minimal 6 paragraf untuk setiap poin pembahasan. Uraikan setiap aspek secara detail, berikan contoh, dan sintesis informasi dari berbagai sumber untuk membangun argumen yang kuat."
        if length_preference == 'Normal':
            length_instruction = "Tulis pembahasan dengan detail yang seimbang, sekitar 2-4 paragraf untuk setiap poin."
        elif length_preference == 'Ringkas':
            length_instruction = "Tulis pembahasan yang ringkas dan padat, sekitar 1-2 paragraf untuk setiap poin."

        prompt_draft = f"""
        Anda adalah seorang penulis akademik ahli dengan standar tertinggi. Tugas Anda adalah menulis konten HANYA untuk satu sub-bab berikut dengan sangat teliti.

        Judul Penelitian Utama: "{research_title}"
        Sub-bab yang Harus Ditulis: "{subchapter.get('sub_bab')}"
        Poin-Poin Kunci yang WAJIB Dibahas: {json.dumps(subchapter.get('poin_pembahasan', []), ensure_ascii=False)}

        Sumber Rujukan yang Tersedia (Gunakan ini sebagai satu-satunya sumber kebenaran):
        {sources_text}

        INSTRUKSI PENULISAN SANGAT PENTING DAN TIDAK BOLEH DILANGGAR:
        1.  **Struktur Tulisan**: Strukturkan jawaban Anda dengan membahas setiap 'poin_pembahasan' secara berurutan. **WAJIB GUNAKAN PENOMORAN ANGKA (1., 2., 3., dst.)** untuk setiap poin di dalam tulisan Anda. Setiap nomor HARUS diikuti dengan penjelasan mendalam.
        2.  **Panjang dan Kedalaman**: {length_instruction}
        3.  **ATURAN SITASI MUTLAK**:
            - SETIAP KLAIM, DEFINISI, ATAU DATA HARUS DIDUKUNG OLEH SITASI. JANGAN PERNAH menulis kalimat atau paragraf tanpa menyertakan setidaknya satu placeholder sitasi [NamaPenulis, Tahun] dari sumber yang relevan.
            - **JANGAN PERNAH MENGGUNAKAN SITASI DARI SUMBER YANG SAMA LEBIH DARI SATU KALI DALAM SATU PARAGRAF.** Gabungkan ide dari sumber yang sama, lalu letakkan SATU sitasi di akhir paragraf tersebut.
            - Jika Anda benar-benar tidak menemukan informasi untuk suatu poin dari daftar sumber yang diberikan, dan HANYA jika demikian, tulis: "Tidak ditemukan pembahasan spesifik mengenai poin ini dalam referensi yang tersedia." JANGAN PERNAH mengarang informasi.
        4.  **Format**: JANGAN tulis judul sub-bab (seperti 'A. Landasan Teori'). Langsung mulai dengan penomoran poin (1., 2., dst.) dan konten paragrafnya dalam format Markdown.
        5.  **Daftar Pustaka**: Buat bagian `### Daftar Pustaka` di akhir. Cantumkan HANYA referensi yang Anda kutip, dan WAJIB sertakan DOI jika tersedia dari daftar sumber.

        Mulai penulisan konten yang detail dan penuh sitasi untuk sub-bab ini sekarang.
        """
        
        draft_response = model.generate_content(prompt_draft)
        generated_text = draft_response.text

        final_text = generated_text
        temp_ref_map = {ref['citation_placeholder'].lower(): ref for ref in processed_references}
        
        placeholders_in_text = re.findall(r'\[([\w\s&.,]+),\s*(\d{4}|n\.d\.)\]', generated_text)
        for author_match, year_match in placeholders_in_text:
            placeholder_key = f"[{author_match}, {year_match}]".lower()
            matched_ref = temp_ref_map.get(placeholder_key)
            if matched_ref:
                in_text_citation = f"({matched_ref['authors_str']}, {matched_ref['year']})"
                final_text = final_text.replace(matched_ref['citation_placeholder'], in_text_citation, 1)

        return jsonify({"generated_text": final_text})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Terjadi kesalahan saat generate konten: {str(e)}"}), 500

# =========================================================================
# API LAINNYA
# =========================================================================

@app.route('/api/export-document', methods=['POST'])
@login_required
def export_document():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'export_doc')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429

    try:
        data = request.get_json()
        html_content = data.get('content')
        export_format = data.get('format')
        title = data.get('title', 'Dokumen-OnThesis')

        if not html_content or not export_format:
            return jsonify({'error': 'Konten dan format ekspor diperlukan.'}), 400

        text_content = html_content.replace('<br>', '\n').replace('</p>', '\n\n').replace('<h3>', '').replace('</h3>', '\n')
        text_content = re.sub('<[^<]+?>', '', text_content) 

        buffer = io.BytesIO()
        
        logo_path = os.path.join(app.static_folder, 'images', 'logo.png')

        if export_format == 'pdf':
            if not PDF_EXPORT_ENABLED:
                return jsonify({'error': 'Fungsi ekspor PDF tidak tersedia di server.'}), 501
            
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            if os.path.exists(logo_path):
                img = Image(logo_path, width=1.5*inch, height=0.5*inch)
                img.hAlign = 'RIGHT'
                story.append(img)
                story.append(Spacer(1, 0.25*inch))

            story.append(Paragraph(title, styles['h1']))
            story.append(Spacer(1, 0.2*inch))

            for line in text_content.split('\n'):
                p = Paragraph(line, styles['BodyText'])
                story.append(p)
                story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            mimetype = 'application/pdf'
            filename = f'{secure_filename(title)}.pdf'

        elif export_format == 'word':
            if not WORD_EXPORT_ENABLED:
                return jsonify({'error': 'Fungsi ekspor Word tidak tersedia di server.'}), 501

            document = Document()
            if os.path.exists(logo_path):
                document.add_picture(logo_path, width=Inches(1.5))
            
            document.add_heading(title, level=1)
            document.add_paragraph(text_content)
            document.save(buffer)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            filename = f'{secure_filename(title)}.docx'
        
        else:
            return jsonify({'error': 'Format tidak didukung.'}), 400

        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype=mimetype)

    except Exception as e:
        print(f"Error saat ekspor dokumen: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/interpret-analysis', methods=['POST'])
@login_required
def interpret_analysis():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429
    try:
        stats_text = request.get_json().get('stats')
        if not stats_text: return jsonify({'error': 'Data statistik tidak boleh kosong.'}), 400
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"Anda adalah seorang analis data. Berdasarkan data statistik berikut:\n---\n{stats_text}\n---\nBerikan interpretasi singkat yang mudah dipahami dalam format markdown."
        response = model.generate_content(prompt)
        return jsonify({'interpretation': response.text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/search-references', methods=['POST'])
@login_required
def api_search_references():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_usage(current_user.id, 'search')
        if not is_allowed: return jsonify({'error': message}), 429
    
    data = request.get_json()
    source = data.get('source')
    query = data.get('query')
    year = data.get('year')
    try:
        if source == 'core':
            core_api_key = os.getenv('CORE_API_KEY')
            if not core_api_key: return jsonify({'error': 'Kunci API CORE tidak dikonfigurasi.'}), 500
            api_url = 'https://api.core.ac.uk/v3/search/works'
            q = f"(title:({query}) OR authors:({query}))"
            if year: q += f" AND yearPublished:{year}"
            params = {'q': q, 'limit': 20}
            headers = {'Authorization': f'Bearer {core_api_key}'}
            response = requests.get(api_url, params=params, headers=headers, timeout=20)
            response.raise_for_status()
            return jsonify(response.json())
        elif source == 'crossref':
            base_url = 'https://api.crossref.org/works'
            params = {'query.bibliographic': query, 'rows': 20}
            if year: params['filter'] = f'from-pub-date:{year}-01-01,until-pub-date:{year}-12-31'
            headers = {'User-Agent': 'OnThesisApp/1.0 (mailto:contact@onthesis.app)'}
            response = requests.get(base_url, params=params, headers=headers, timeout=20)
            response.raise_for_status()
            api_data = response.json()
            results = api_data.get('message', {}).get('items', [])
            return jsonify({'results': results})
        else:
            return jsonify({'error': 'Sumber tidak valid.'}), 400
    except Exception as e:
        return jsonify({'error': f'Terjadi kesalahan saat mencari referensi: {e}'}), 500

@app.route('/paraphrase', methods=['POST'])
@login_required
def paraphrase_text():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_usage(current_user.id, 'paraphrase')
        if not is_allowed: return jsonify({'error': message}), 429
    
    try:
        data = request.get_json()
        text = data.get('text')
        intensity_level = data.get('intensity', '2')

        if not text: return jsonify({'error': 'Teks tidak boleh kosong.'}), 400
        
        intensity_map = {
            '1': 'sedikit mengubah struktur kalimat dan beberapa kata kunci',
            '2': 'mengubah struktur kalimat secara signifikan dan mengganti banyak sinonim',
            '3': 'menulis ulang sepenuhnya dengan gaya yang sangat berbeda namun tetap mempertahankan ide inti'
        }
        instruction = intensity_map.get(intensity_level, intensity_map['2'])
        
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        Anda adalah seorang ahli parafrase untuk tulisan akademis.
        Tugas Anda adalah memparafrasekan teks berikut dengan gaya penulisan untuk skripsi.
        Instruksi spesifik: {instruction}.
        Pastikan makna asli tetap terjaga.

        Teks Asli:
        ---
        {text}
        ---

        Hasil Parafrase:
        """
        response = model.generate_content(prompt)
        return jsonify({'paraphrased_text': response.text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/chat', methods=['POST'])
@login_required
def chat_with_ai():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_usage(current_user.id, 'chat')
        if not is_allowed: return jsonify({'error': message}), 429

    try:
        message = request.get_json().get('message')
        if not message: return jsonify({'error': 'Pesan tidak boleh kosong.'}), 400
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"Anda adalah asisten AI bernama OnThesis. Jawab pertanyaan mahasiswa ini seputar skripsi dengan ramah dan membantu: {message}"
        response = model.generate_content(prompt)
        return jsonify({'reply': response.text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analyze-document', methods=['POST'])
@login_required
def analyze_document():
    if 'document' not in request.files:
        return jsonify({'error': 'Tidak ada file yang diunggah.'}), 400
    file = request.files['document']
    if file.filename == '':
        return jsonify({'error': 'Nama file kosong.'}), 400
    try:
        filename = secure_filename(file.filename).lower()
        content = ""
        if filename.endswith('.pdf'):
            content = read_pdf(file.stream)
        elif filename.endswith('.docx'):
            content = read_docx(file.stream)
        else:
            return jsonify({'error': 'Format file tidak didukung. Harap unggah PDF atau DOCX.'}), 400
        if not content.strip():
            return jsonify({'error': 'Tidak ada teks yang dapat diekstrak dari file ini.'}), 400
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        Dari teks dokumen akademis berikut, identifikasi informasi sitasi untuk dokumen itu sendiri.
        Ekstrak penulis utama, judul utama, tahun publikasi, dan nama jurnal atau konferensi tempat dokumen itu diterbitkan.
        Berikan hasilnya sebagai array JSON yang hanya berisi SATU objek dengan kunci: "title", "author", "year", dan "journal".

        Teks Dokumen (ambil dari bagian awal untuk efisiensi):
        ---
        {content[:8000]} 
        ---
        """
        response = model.generate_content(prompt)
        clean_json_string = re.sub(r'```json\s*|\s*```', '', response.text.strip(), flags=re.DOTALL)
        if not clean_json_string.strip().startswith('['):
            clean_json_string = f"[{clean_json_string}]"
        references = json.loads(clean_json_string)
        return jsonify({'references': references})
    except json.JSONDecodeError:
        return jsonify({'error': 'AI tidak dapat memformat informasi sitasi dengan benar. Coba lagi.'}), 500
    except Exception as e:
        print(f"Error saat menganalisis dokumen: {e}")
        return jsonify({'error': f'Terjadi kesalahan internal: {str(e)}'}), 500

@app.route('/api/normality', methods=['POST'])
@login_required
def api_normality():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429
    try:
        data = request.get_json()
        values = data.get('values')
        if not values or not isinstance(values, list):
            return jsonify({'error': 'Data angka diperlukan dalam array'}), 400
        
        s = pd.Series(values, dtype=float).dropna()
        
        if len(s) < 3:
            return jsonify({'error': 'Minimal 3 data untuk uji normalitas'}), 400
        
        n = len(s)
        mean = s.mean()
        sd = s.std(ddof=1)

        shapiro_stat, shapiro_p = stats.shapiro(s)
        ks_stat, ks_p = stats.kstest((s - mean) / sd, 'norm')

        shapiro_p_rounded = round(shapiro_p, 3)
        ks_p_rounded = round(ks_p, 3)

        if shapiro_p > 0.05:
            conclusion = "berdistribusi normal"
        else:
            conclusion = "tidak berdistribusi normal"
        
        summary = f"Hasil uji Shapiro-Wilk menunjukkan nilai signifikansi p = {shapiro_p_rounded}. Karena nilai p > 0.05, dapat disimpulkan bahwa data {conclusion}."

        df_table = pd.DataFrame({
            "test": ["Shapiro-Wilk", "Kolmogorov-Smirnov"],
            "statistic": [shapiro_stat, ks_stat],
            "df": [n, n],
            "p": [shapiro_p_rounded, ks_p_rounded]
        })
        table_json = df_table.to_dict(orient='records')

        return jsonify({
            "summary": summary, 
            "mean": round(mean, 3), 
            "std_dev": round(sd, 3), 
            "n": n, 
            "table": table_json
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/levene', methods=['POST'])
@login_required
def api_levene():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429
    try:
        data = request.get_json()
        groups = data.get('groups')
        if not groups or not isinstance(groups, list) or len(groups) < 2:
            return jsonify({'error': 'Minimal 2 grup data diperlukan'}), 400

        cleaned_groups = [np.array(g, dtype=float)[~np.isnan(g)] for g in groups]
        group_sizes = [len(g) for g in cleaned_groups]
        if any(size < 2 for size in group_sizes):
            return jsonify({'error': 'Setiap grup minimal punya 2 data'}), 400

        stat, p_value = stats.levene(*cleaned_groups)
        
        p_string = f"{p_value:.3f}"
        comparison = "> 0.05" if p_value > 0.05 else "<= 0.05"
        conclusion = "homogen" if p_value > 0.05 else "tidak homogen"
        summary = f"Hasil Levene’s Test menunjukkan nilai Sig. = {p_string} ({comparison}), sehingga dapat disimpulkan varians data antar kelompok adalah {conclusion}."

        df_table = pd.DataFrame({
            "test": ["Levene’s Test"],
            "statistic": [stat],
            "df1": [len(cleaned_groups) - 1],
            "df2": [sum(group_sizes) - len(cleaned_groups)],
            "p": [p_value]
        })
        df_table['statistic'] = df_table['statistic'].round(4)
        df_table['p'] = df_table['p'].round(4)
        table_json = df_table.to_dict(orient='records')

        return jsonify({
            "summary": summary,
            "n_per_group": group_sizes,
            "table": table_json
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/bartlett', methods=['POST'])
@login_required
def api_bartlett():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429
    try:
        data = request.get_json()
        groups = data.get('groups')
        if not groups or not isinstance(groups, list) or len(groups) < 2:
            return jsonify({'error': 'Minimal 2 grup data diperlukan'}), 400

        cleaned_groups = [np.array(g, dtype=float)[~np.isnan(g)] for g in groups]
        group_sizes = [len(g) for g in cleaned_groups]
        if any(size < 2 for size in group_sizes):
            return jsonify({'error': 'Setiap grup minimal punya 2 data'}), 400

        stat, p_value = stats.bartlett(*cleaned_groups)

        p_string = f"{p_value:.3f}"
        comparison = "> 0.05" if p_value > 0.05 else "<= 0.05"
        conclusion = "homogen" if p_value > 0.05 else "tidak homogen"
        summary = f"Hasil Bartlett's Test menunjukkan nilai Sig. = {p_string} ({comparison}), sehingga dapat disimpulkan varians data antar kelompok adalah {conclusion}."

        df_table = pd.DataFrame({
            "test": ["Bartlett’s Test"],
            "statistic": [stat],
            "df1": [len(cleaned_groups) - 1],
            "df2": [None],
            "p": [p_value]
        })
        df_table['statistic'] = df_table['statistic'].round(4)
        df_table['p'] = df_table['p'].round(4)
        table_json = df_table.to_dict(orient='records')

        return jsonify({
            "summary": summary,
            "n_per_group": group_sizes,
            "table": table_json
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
        
@app.route('/api/descriptive-analysis', methods=['POST'])
@login_required
def api_descriptive_analysis():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429

    try:
        if not request.is_json:
            return jsonify({'error': 'Format request tidak valid, harus JSON.'}), 400

        data = request.get_json()
        if not data or not any(data.values()):
             return jsonify({'error': 'Tidak ada data yang dikirim untuk dianalisis.'}), 400
        
        df = pd.DataFrame.from_dict(data, orient='index').transpose()
        
        for col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')

        if df.empty:
            return jsonify({'error': 'Data kosong setelah diproses.'}), 400

        all_cols = df.columns.tolist()
        results = {}
        plots = {}
        
        for col in all_cols:
            series = df[col].dropna()
            if len(series) < 2: continue

            results[col] = {
                'n': int(len(series)),
                'mean': float(series.mean()),
                'median': float(series.median()),
                'mode': series.mode().tolist() if not series.mode().empty else ['N/A'],
                'std': float(series.std()),
                'variance': float(series.var()),
                'range': float(series.max() - series.min()),
                'min': float(series.min()),
                'max': float(series.max()),
            }

            sns.set_style("whitegrid")
            
            fig_hist, ax_hist = plt.subplots()
            sns.histplot(series, kde=True, ax=ax_hist, color='#0284c7')
            ax_hist.set_title(f'Histogram - {col}')
            plots[f'{col}_histogram'] = create_plot_as_base64(fig_hist)

            fig_box, ax_box = plt.subplots()
            sns.boxplot(x=series, ax=ax_box, color='#0284c7')
            ax_box.set_title(f'Box Plot - {col}')
            plots[f'{col}_boxplot'] = create_plot_as_base64(fig_box)

            try:
                if series.nunique() > 1 and series.nunique() <= 10:
                    pie_data = series.value_counts()
                    fig_pie, ax_pie = plt.subplots()
                    ax_pie.pie(pie_data, labels=pie_data.index.astype(str), autopct='%1.1f%%', startangle=90)
                    ax_pie.axis('equal')
                    ax_pie.set_title(f'Distribusi Proporsi - {col}')
                    plots[f'{col}_piechart'] = create_plot_as_base64(fig_pie)
                else:
                    plots[f'{col}_piechart'] = None
            except Exception as pie_e:
                print(f"Tidak dapat membuat pie chart untuk {col}: {pie_e}")
                plots[f'{col}_piechart'] = None
        return jsonify({
            'columns': all_cols,
            'results': results,
            'plots': plots
        })

    except Exception as e:
        print(f"Error in descriptive_analysis API: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Terjadi kesalahan internal: {str(e)}'}), 500


@app.route('/api/get-usage-status')
@login_required
def get_usage_status():
    if current_user.is_pro:
        return jsonify({'status': 'pro', 'message': 'Akses Penuh Tanpa Batas'})
    LIMITS = {'paraphrase': 5, 'chat': 10, 'search': 5, 'citation': 15, 'writing_assistant': 3, 'data_analysis': 3, 'export_doc': 1, 'generate_theory': 2}
    user_ref = db.collection('users').document(current_user.id)
    user_doc = user_ref.get()
    if not user_doc.exists: return jsonify({'error': 'User not found'}), 404
    usage_data = user_doc.to_dict().get('usage_limits', {})
    today_str = date.today().isoformat()
    if usage_data.get('last_reset_date') != today_str:
        usage_data['paraphrase_count'] = 0
        usage_data['chat_count'] = 0
        usage_data['search_count'] = 0
    return jsonify({
        'status': 'free',
        'paraphrase_remaining': LIMITS['paraphrase'] - usage_data.get('paraphrase_count', 0),
        'chat_remaining': LIMITS['chat'] - usage_data.get('chat_count', 0),
        'search_remaining': LIMITS['search'] - usage_data.get('search_count', 0),
        'citation_remaining': LIMITS['citation'] - usage_data.get('citation_count', 0),
        'writing_assistant_remaining': LIMITS['writing_assistant'] - usage_data.get('writing_assistant_count', 0),
        'data_analysis_remaining': LIMITS['data_analysis'] - usage_data.get('data_analysis_count', 0),
        'export_doc_remaining': LIMITS['export_doc'] - usage_data.get('export_doc_count', 0),
        'generate_theory_remaining': LIMITS['generate_theory'] - usage_data.get('generate_theory_count', 0),
        'limits': LIMITS
    })

@app.route('/api/submit-feedback', methods=['POST'])
@login_required
def submit_feedback():
    try:
        data = request.get_json()
        message = data.get('message')
        category = data.get('category')
        page_url = data.get('pageUrl')
        if not message or not category:
            return jsonify({'status': 'error', 'message': 'Pesan dan kategori tidak boleh kosong.'}), 400
        feedback_doc = {
            'userId': current_user.id, 'userEmail': current_user.email,
            'message': message, 'category': category, 'pageUrl': page_url,
            'timestamp': firestore.SERVER_TIMESTAMP, 'status': 'new'
        }
        db.collection('feedback').add(feedback_doc)
        return jsonify({'status': 'success', 'message': 'Terima kasih atas masukan Anda!'})
    except Exception as e:
        print(f"Error saat menyimpan feedback: {e}")
        return jsonify({'status': 'error', 'message': 'Terjadi kesalahan di server.'}), 500

@app.route('/api/create-transaction', methods=['POST'])
@login_required
def create_transaction():
    try:
        if not midtrans_snap:
            return jsonify({'status': 'error', 'message': 'Layanan pembayaran tidak terkonfigurasi.'}), 503
        data = request.get_json()
        plan = data.get('plan')
        amount = data.get('amount')
        if not plan or not amount:
            return jsonify({'status': 'error', 'message': 'Detail paket tidak lengkap.'}), 400
        order_id = f"ONTESIS-PRO-{current_user.id}-{plan}-{int(time.time())}"
        transaction_details = {"order_id": order_id, "gross_amount": int(amount)}
        customer_details = {"first_name": current_user.displayName, "email": current_user.email}
        transaction = midtrans_snap.create_transaction({
            "transaction_details": transaction_details,
            "customer_details": customer_details
        })
        return jsonify({'status': 'success', 'token': transaction['token']})
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'Gagal membuat transaksi: {e}'}), 500

@app.route('/api/payment-notification', methods=['POST'])
def payment_notification():
    try:
        notification_json = request.get_json()
        order_id = notification_json['order_id']
        transaction_status = notification_json['transaction_status']
        fraud_status = notification_json.get('fraud_status')
        if transaction_status == 'settlement' and fraud_status == 'accept':
            parts = order_id.split('-')
            if len(parts) >= 4 and parts[0] == 'ONTESIS' and parts[1] == 'PRO':
                user_id = parts[2]
                plan = parts[3]
                now = datetime.now()
                expiry_date = None
                if plan == 'weekly':
                    expiry_date = now + timedelta(days=7)
                elif plan == 'monthly':
                    expiry_date = now + timedelta(days=30)
                elif plan == 'yearly':
                    expiry_date = now + timedelta(days=365)
                if expiry_date:
                    db.collection('users').document(user_id).update({
                        'proExpiryDate': expiry_date,
                        'lastSubscriptionPlan': plan
                    })
                    print(f"Sukses: Pengguna {user_id} telah upgrade ke paket {plan}.")
        return jsonify({'status': 'ok'}), 200
    except Exception as e:
        print(f"Error saat menangani notifikasi pembayaran: {e}")
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500

# --- API UNTUK UJI T ---
@app.route('/api/independent-ttest', methods=['POST'])
@login_required
def api_independent_ttest():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429
    
    try:
        data = request.get_json()
        groups = data.get('groups')
        confidence_level = float(data.get('confidence_level', 95)) / 100.0

        if not groups or len(groups) != 2:
            return jsonify({'error': 'Dibutuhkan tepat dua grup data.'}), 400

        group1 = np.array(groups[0], dtype=float)
        group2 = np.array(groups[1], dtype=float)

        if len(group1) < 3 or len(group2) < 3:
            return jsonify({'error': 'Setiap grup harus memiliki minimal 3 data poin.'}), 400

        stats1 = {'N': len(group1), 'mean': np.mean(group1), 'std': np.std(group1, ddof=1)}
        stats2 = {'N': len(group2), 'mean': np.mean(group2), 'std': np.std(group2, ddof=1)}
        
        levene_stat, levene_p = stats.levene(group1, group2)

        t_stat_equal, p_equal = stats.ttest_ind(group1, group2, equal_var=True)
        t_stat_unequal, p_unequal = stats.ttest_ind(group1, group2, equal_var=False)
        
        df_equal = len(group1) + len(group2) - 2
        v1 = np.var(group1, ddof=1)
        v2 = np.var(group2, ddof=1)
        n1 = len(group1)
        n2 = len(group2)
        df_unequal = (v1/n1 + v2/n2)**2 / ( (v1/n1)**2/(n1-1) + (v2/n2)**2/(n2-1) )

        mean_diff = np.mean(group1) - np.mean(group2)
        
        se_diff_equal = np.sqrt( ( (n1-1)*v1 + (n2-1)*v2 ) / df_equal * (1/n1 + 1/n2) )
        ci_equal = stats.t.interval(confidence_level, df_equal, loc=mean_diff, scale=se_diff_equal)

        se_diff_unequal = np.sqrt(v1/n1 + v2/n2)
        ci_unequal = stats.t.interval(confidence_level, df_unequal, loc=mean_diff, scale=se_diff_unequal)
        
        p_to_check = levene_p > 0.05 and p_equal or p_unequal
        conclusion = "terdapat perbedaan rata-rata yang signifikan" if p_to_check < 0.05 else "tidak terdapat perbedaan rata-rata yang signifikan"
        summary = f"Berdasarkan hasil uji T (p = {p_to_check:.3f}), dapat disimpulkan bahwa {conclusion} antara kedua kelompok."

        result = {
            'summary': summary,
            'group_stats': [
                {'group': 'Grup 1', **stats1, 'ci_lower': stats.t.interval(confidence_level, len(group1)-1, loc=np.mean(group1), scale=stats.sem(group1))[0], 'ci_upper': stats.t.interval(confidence_level, len(group1)-1, loc=np.mean(group1), scale=stats.sem(group1))[1]},
                {'group': 'Grup 2', **stats2, 'ci_lower': stats.t.interval(confidence_level, len(group2)-1, loc=np.mean(group2), scale=stats.sem(group2))[0], 'ci_upper': stats.t.interval(confidence_level, len(group2)-1, loc=np.mean(group2), scale=stats.sem(group2))[1]}
            ],
            'independent_test': {
                'levene': {'F': levene_stat, 'p': levene_p},
                'ttest_equal_variances': {'t': t_stat_equal, 'df': df_equal, 'p': p_equal, 'mean_diff': mean_diff, 'ci_lower': ci_equal[0], 'ci_upper': ci_equal[1]},
                'ttest_unequal_variances': {'t': t_stat_unequal, 'df': df_unequal, 'p': p_unequal, 'mean_diff': mean_diff, 'ci_lower': ci_unequal[0], 'ci_upper': ci_unequal[1]}
            }
        }
        return jsonify(sanitize_nan(result))
    except Exception as e:
        print(f"Error di api_independent_ttest: {e}")
        return jsonify({'error': 'Terjadi kesalahan saat memproses data. Pastikan format data benar.'}), 500


@app.route('/api/paired-ttest', methods=['POST'])
@login_required
def api_paired_ttest():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed:
            if message == "UPGRADE_REQUIRED":
                return jsonify({'error': "Batas percobaan tercapai.", 'redirect': url_for('upgrade_page')}), 429
            return jsonify({'error': message}), 429
    try:
        data = request.get_json()
        pairs = data.get('pairs')
        confidence_level = float(data.get('confidence_level', 95)) / 100.0

        if not pairs or len(pairs) != 2:
            return jsonify({'error': 'Dibutuhkan tepat dua set data berpasangan.'}), 400

        pair1 = np.array(pairs[0], dtype=float)
        pair2 = np.array(pairs[1], dtype=float)

        if len(pair1) != len(pair2) or len(pair1) < 3:
            return jsonify({'error': 'Kedua set data harus memiliki jumlah yang sama dan minimal 3 data poin.'}), 400

        stats1 = {'N': len(pair1), 'mean': np.mean(pair1), 'std': np.std(pair1, ddof=1)}
        stats2 = {'N': len(pair2), 'mean': np.mean(pair2), 'std': np.std(pair2, ddof=1)}
        
        corr_r, corr_p = stats.pearsonr(pair1, pair2)

        t_stat, p_value = stats.ttest_rel(pair1, pair2)
        
        diff = pair1 - pair2
        mean_diff = np.mean(diff)
        std_diff = np.std(diff, ddof=1)
        df = len(pair1) - 1
        
        if std_diff > 0:
            ci = stats.t.interval(confidence_level, df, loc=mean_diff, scale=stats.sem(diff))
        else:
            ci = (mean_diff, mean_diff)

        conclusion = "terdapat perbedaan rata-rata yang signifikan" if p_value < 0.05 else "tidak terdapat perbedaan rata-rata yang signifikan"
        summary = f"Berdasarkan hasil uji T berpasangan (p = {p_value:.3f}), dapat disimpulkan bahwa {conclusion} antara kedua pengukuran."

        result = {
            'summary': summary,
            'paired_stats': [
                {'variable': 'Variabel 1', **stats1},
                {'variable': 'Variabel 2', **stats2}
            ],
            'paired_correlation': {
                'pair': 'Variabel 1 & Variabel 2',
                'r': corr_r,
                'p': corr_p
            },
            'paired_test': {
                'pair': 'Variabel 1 - Variabel 2',
                'mean_diff': mean_diff,
                'std_diff': std_diff,
                't': t_stat,
                'df': df,
                'p': p_value,
                'ci_lower': ci[0],
                'ci_upper': ci[1]
            }
        }
        return jsonify(sanitize_nan(result))
    except Exception as e:
        print(f"Error di api_paired_ttest: {e}")
        return jsonify({'error': 'Terjadi kesalahan saat memproses data. Pastikan format data benar.'}), 500


# ========================================================================
# FUNGSI-FUNGSI ANALISIS ANOVA (DIPERBARUI DENGAN NAMA FUNGSI YANG BENAR)
# ========================================================================

def _perform_oneway_anova_analysis(df, dependent_var, independent_var):
    """Fungsi helper khusus untuk One-Way ANOVA."""
    df_cleaned = df[[dependent_var, independent_var]].copy()
    df_cleaned[dependent_var] = pd.to_numeric(df_cleaned[dependent_var], errors='coerce')
    df_cleaned.dropna(inplace=True)

    if df_cleaned[independent_var].nunique() < 2:
        raise ValueError(f"Analisis ANOVA membutuhkan minimal 2 kelompok data. Kolom grup ('{independent_var}') Anda hanya memiliki {df_cleaned[independent_var].nunique()} kelompok unik.")
    
    group_counts = df_cleaned.groupby(independent_var)[dependent_var].count()
    if (group_counts < 2).any():
        invalid_groups = group_counts[group_counts < 2].index.tolist()
        raise ValueError(f"Setiap kelompok harus memiliki minimal 2 data poin yang valid. Kelompok berikut tidak memenuhi syarat: {', '.join(invalid_groups)}.")

    df_cleaned[independent_var] = df_cleaned[independent_var].astype('category')

    normality_results = pg.normality(data=df_cleaned, dv=dependent_var, group=independent_var)
    is_all_normal = bool(all(normality_results['normal']))
    
    homogeneity_result = pg.homoscedasticity(data=df_cleaned, dv=dependent_var, group=independent_var, method='levene')
    is_homogeneous = bool(homogeneity_result['equal_var'].iloc[0]) if not homogeneity_result.empty else False

    if is_all_normal:
        analysis_type = "One-Way ANOVA"
        aov = pg.anova(data=df_cleaned, dv=dependent_var, between=independent_var, detailed=True)
        main_test_results = json.loads(aov.round(4).to_json(orient='records'))[0]
        p_value = main_test_results['p-unc']
        f_stat = main_test_results['F']
        df_between = main_test_results['DF']
        df_within = aov.loc[1, 'DF']

        if p_value < 0.05:
            post_hoc = pg.pairwise_tukey(data=df_cleaned, dv=dependent_var, between=independent_var) if is_homogeneous else pg.pairwise_gameshowell(data=df_cleaned, dv=dependent_var, between=independent_var)
            post_hoc_results = json.loads(post_hoc.round(4).to_json(orient='records'))
            summary_indonesia = f"Terdapat perbedaan yang signifikan secara statistik pada rata-rata '{dependent_var}' antar kelompok '{independent_var}', F({df_between}, {df_within}) = {f_stat:.2f}, p = {p_value:.3f}."
            summary_apa = f"A one-way ANOVA revealed a significant effect of {independent_var} on {dependent_var}, F({df_between}, {df_within}) = {f_stat:.2f}, p < .05."
        else:
            post_hoc_results = None
            summary_indonesia = f"Tidak ditemukan perbedaan yang signifikan secara statistik pada rata-rata '{dependent_var}' antar kelompok '{independent_var}', F({df_between}, {df_within}) = {f_stat:.2f}, p = {p_value:.3f}."
            summary_apa = f"A one-way ANOVA did not reveal a significant effect of {independent_var} on {dependent_var}, F({df_between}, {df_within}) = {f_stat:.2f}, p > .05."
    else:
        analysis_type = "Kruskal-Wallis H Test"
        kruskal = pg.kruskal(data=df_cleaned, dv=dependent_var, between=independent_var)
        main_test_results = json.loads(kruskal.round(4).to_json(orient='records'))[0]
        p_value = main_test_results['p-unc']
        h_stat = main_test_results['H']
        df_kruskal = main_test_results['ddof1']
        post_hoc_results = None
        summary_indonesia = f"Hasil uji Kruskal-Wallis menunjukkan tidak ada perbedaan peringkat (rank) yang signifikan secara statistik pada '{dependent_var}' antar kelompok '{independent_var}', H({df_kruskal}) = {h_stat:.2f}, p = {p_value:.3f}."
        summary_apa = f"A Kruskal-Wallis H test showed no statistically significant difference in ranks for {dependent_var} across {independent_var} groups, H({df_kruskal}) = {h_stat:.2f}, p > .05."


    descriptive_stats = df_cleaned.groupby(independent_var)[dependent_var].describe().round(3)
    
    plt.figure(figsize=(10, 6))
    sns.boxplot(x=independent_var, y=dependent_var, data=df_cleaned, palette="pastel")
    sns.stripplot(x=independent_var, y=dependent_var, data=df_cleaned, color=".25", size=4)
    plt.title(f'Distribusi {dependent_var} Berdasarkan {independent_var}', fontsize=16)
    plot_base64 = create_plot_as_base64(plt.gcf())

    return {
        'success': True, 'analysis_type': analysis_type,
        'prerequisites': { 'normality': json.loads(normality_results.round(4).to_json(orient='records')), 'homogeneity': json.loads(homogeneity_result.round(4).to_json(orient='records')), 'is_all_normal': is_all_normal, 'is_homogeneous': is_homogeneous },
        'descriptive_stats': json.loads(descriptive_stats.reset_index().to_json(orient='records')),
        'main_test_results': main_test_results, 'post_hoc_results': post_hoc_results,
        'summary': {'apa': summary_apa, 'indonesia': summary_indonesia},
        'plot': plot_base64
    }

def _perform_twoway_anova_analysis(df, dependent_var, independent_vars):
    """Fungsi helper BARU khusus untuk Two-Way ANOVA."""
    iv1, iv2 = independent_vars[0], independent_vars[1]
    
    df_cleaned = df[[dependent_var, iv1, iv2]].copy()
    df_cleaned[dependent_var] = pd.to_numeric(df_cleaned[dependent_var], errors='coerce')
    df_cleaned.dropna(inplace=True)

    if df_cleaned[iv1].nunique() < 2 or df_cleaned[iv2].nunique() < 2:
        raise ValueError("Setiap variabel independen pada Two-Way ANOVA harus memiliki minimal 2 level/kategori.")

    interaction_col = f"{iv1}_{iv2}"
    df_cleaned[interaction_col] = df_cleaned[iv1].astype(str) + "_" + df_cleaned[iv2].astype(str)

    normality_results = pg.normality(data=df_cleaned, dv=dependent_var, group=iv1)
    is_all_normal = bool(all(normality_results['normal']))
    
    homogeneity_result = pg.homoscedasticity(data=df_cleaned, dv=dependent_var, group=interaction_col, method='levene')
    is_homogeneous = bool(homogeneity_result['equal_var'].iloc[0]) if not homogeneity_result.empty else False

    df_cleaned.drop(columns=[interaction_col], inplace=True)

    aov = pg.anova(data=df_cleaned, dv=dependent_var, between=[iv1, iv2], detailed=True)
    main_test_results = json.loads(aov.round(4).to_json(orient='records'))

    p_interaction = aov[aov['Source'] == f'{iv1} * {iv2}']['p-unc'].iloc[0]
    summary_indonesia = f"Hasil Two-Way ANOVA menunjukkan "
    summary_apa = "A two-way ANOVA was conducted to examine the effects of " + iv1 + " and " + iv2 + " on " + dependent_var + ". "
    
    if p_interaction < 0.05:
        summary_indonesia += f"terdapat efek interaksi yang signifikan antara '{iv1}' dan '{iv2}' terhadap '{dependent_var}' (p < .05)."
        summary_apa += "There was a statistically significant interaction between the effects of " + iv1 + " and " + iv2 + "."
    else:
        summary_indonesia += f"tidak terdapat efek interaksi yang signifikan (p > .05). "
        summary_apa += "There was not a statistically significant interaction between the effects of " + iv1 + " and " + iv2 + ". "
        p_iv1 = aov[aov['Source'] == iv1]['p-unc'].iloc[0]
        p_iv2 = aov[aov['Source'] == iv2]['p-unc'].iloc[0]
        if p_iv1 < 0.05: summary_indonesia += f"Terdapat efek utama yang signifikan dari '{iv1}'. "
        if p_iv2 < 0.05: summary_indonesia += f"Terdapat efek utama yang signifikan dari '{iv2}'. "

    plt.figure(figsize=(10, 6))
    sns.pointplot(data=df_cleaned, x=iv1, y=dependent_var, hue=iv2, dodge=True, errorbar='se', capsize=.1)
    plt.title(f'Grafik Interaksi {iv1} dan {iv2} terhadap {dependent_var}', fontsize=16)
    plot_base64 = create_plot_as_base64(plt.gcf())

    return {
        'success': True, 'analysis_type': 'Two-Way ANOVA',
        'prerequisites': { 'normality': json.loads(normality_results.round(4).to_json(orient='records')), 'homogeneity': json.loads(homogeneity_result.round(4).to_json(orient='records')), 'is_all_normal': is_all_normal, 'is_homogeneous': is_homogeneous },
        'descriptive_stats': json.loads(df_cleaned.groupby([iv1, iv2])[dependent_var].describe().round(3).reset_index().to_json(orient='records')),
        'main_test_results': main_test_results, 'post_hoc_results': None,
        'summary': {'apa': summary_apa, 'indonesia': summary_indonesia},
        'plot': plot_base64
    }


# ========================================================================
# BAGIAN 2: DUA RUTE API YANG DIPERBARUI
# Tempelkan dua fungsi ini untuk menggantikan rute api_anova_test dan
# api_manual_anova_test yang lama.
# ========================================================================
@app.route('/api/anova_test', methods=['POST'], endpoint='api_anova_test_file')
@login_required
def api_anova_test_file():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed: return jsonify({'success': False, 'message': "Batas percobaan tercapai. Upgrade ke PRO."}), 429
            
    if 'file' not in request.files: return jsonify({'success': False, 'message': 'File tidak ditemukan.'}), 400

    try:
        file = request.files['file']
        filename = secure_filename(file.filename)
        if filename.endswith('.csv'): df = pd.read_csv(file)
        elif filename.endswith(('.xls', '.xlsx')): df = pd.read_excel(file)
        else: return jsonify({'success': False, 'message': 'Format file tidak didukung.'}), 400
        
        anova_type = request.form.get('anova_type')
        dependent_var = request.form.get('dependent')
        independent_var1 = request.form.get('independent1')
        independent_var2 = request.form.get('independent2')

        if anova_type == 'one_way':
            result = _perform_oneway_anova_analysis(df, dependent_var, independent_var1)
        elif anova_type == 'two_way':
            if not independent_var2: return jsonify({'success': False, 'message': 'Variabel grup kedua diperlukan untuk Two-Way ANOVA.'}), 400
            result = _perform_twoway_anova_analysis(df, dependent_var, [independent_var1, independent_var2])
        else:
            return jsonify({'success': False, 'message': 'Jenis ANOVA tidak valid.'}), 400
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Terjadi kesalahan saat analisis: {str(e)}'}), 500

@app.route('/api/manual_anova_test', methods=['POST'], endpoint='api_anova_test_manual')
@login_required
def api_manual_anova_test():
    if not current_user.is_pro:
        is_allowed, message = check_and_update_pro_trial(current_user.id, 'data_analysis')
        if not is_allowed: return jsonify({'success': False, 'message': "Batas percobaan tercapai. Upgrade ke PRO."}), 429

    try:
        data = request.get_json()
        anova_type = data.get('anova_type')

        if anova_type == 'one_way':
            groups_data = data.get('groups', [])
            group_names = data.get('group_names', [])
            all_values = []
            for i, group_vals in enumerate(groups_data):
                for val in group_vals:
                    all_values.append({'Nilai': val, 'Kelompok': group_names[i]})
            df = pd.DataFrame(all_values)
            # PERBAIKAN: Memastikan nama fungsi yang dipanggil sudah benar
            result = _perform_oneway_anova_analysis(df, 'Nilai', 'Kelompok')
        elif anova_type == 'two_way':
            table_data = data.get('data', [])
            if not table_data:
                return jsonify({'success': False, 'message': 'Data tabel manual tidak boleh kosong.'}), 400
            df = pd.DataFrame(table_data)
            col_names = list(table_data[0].keys())
            result = _perform_twoway_anova_analysis(df, col_names[0], [col_names[1], col_names[2]])
        else:
            return jsonify({'success': False, 'message': 'Jenis ANOVA tidak valid.'}), 400

        return jsonify(result)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Gagal memproses data manual: {str(e)}'}), 500
